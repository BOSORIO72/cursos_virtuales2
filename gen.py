from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def T(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def C(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def P(text):
    doc.add_paragraph(text)

def B(text):
    doc.add_paragraph(text, style='List Bullet')

def E(tt, code, expl):
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run(">> EJEMPLO: " + tt)
    r.bold = True; r.font.color.rgb = RGBColor(0x00, 0x70, 0x70); r.font.size = Pt(11)
    C(code)
    P("=> " + expl)

def N(text):
    p = doc.add_paragraph()
    r = p.add_run("NOTA: " + text)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(10)

# PORTADA
doc.add_paragraph(""); doc.add_paragraph("")
T("SISTEMA DE CURSOS VIRTUALES", 0)
T("Guia completa del codigo: de principio a fin", 2)
P("")
P("Una explicacion facil de entender con ejemplos practicos")
P("")
p = doc.add_paragraph(); r = p.add_run("Hecho con Laravel 12 | PHP 8.2 | SQLite | Tailwind CSS")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

T("1. Que hace esta aplicacion?", 1)
P("Este proyecto es un panel de administracion para gestionar cursos virtuales. Piensa en el como el sistema que usa una academia online para:")
B("Crear y administrar cursos (con precio gratuito o de pago)")
B("Registrar estudiantes")
B("Inscribir estudiantes en cursos (respetando los cupos disponibles)")
B("Organizar el contenido de cada curso en modulos (algunos gratuitos, otros premium)")
B("Gestionar solicitudes de acceso a cursos de pago")
B("Tener diferentes tipos de usuarios: administrador, editor e invitado")
P("")
P("Todo el sistema esta en espanol y construido con Laravel 12, el framework de PHP mas popular.")
E("Flujo completo del sistema","1. Un profesor crea un curso con 30 cupos y precio S/50\n2. Agrega 5 modulos: 2 gratuitos y 3 premium\n3. Un estudiante se inscribe al curso\n4. El estudiante ve los 2 modulos gratuitos\n5. Para ver los premium, solicita acceso\n6. El administrador aprueba la solicitud\n7. El estudiante ve todos los modulos","Ciclo de vida completo de un estudiante en el sistema.")
doc.add_page_break()

T("2. Tecnologias utilizadas", 1)
t = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
t.cell(0,0).text="Tecnologia"; t.cell(0,1).text="Para que se usa?"
t.cell(1,0).text="Laravel 12"; t.cell(1,1).text="Maneja rutas, controladores, BD, autenticacion"
t.cell(2,0).text="PHP 8.2"; t.cell(2,1).text="Lenguaje de programacion principal"
t.cell(3,0).text="SQLite"; t.cell(3,1).text="Base de datos en un solo archivo"
t.cell(4,0).text="Blade"; t.cell(4,1).text="Sistema de plantillas HTML de Laravel"
t.cell(5,0).text="Tailwind CSS"; t.cell(5,1).text="Framework CSS para dar estilo"
t.cell(6,0).text="Vite"; t.cell(6,1).text="Compila CSS y JavaScript"
print("OK up to section 2")
doc.save(r"C:\Users\brian\OneDrive\Desktop\cursos_virtuales2\explicacion_codigo.docx")
print("Saved OK")
