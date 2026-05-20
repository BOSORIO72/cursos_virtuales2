# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

def T(text, lv=1):
    h = doc.add_heading(text, level=lv)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def C(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def P(text):
    doc.add_paragraph(text)

def B(text):
    doc.add_paragraph(text, style='List Bullet')

def E(tt, code, expl):
    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('>> EJEMPLO: ' + tt)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
    r.font.size = Pt(11)
    C(code)
    P('=> ' + expl)

def N(text):
    p = doc.add_paragraph()
    r = p.add_run('NOTA: ' + text)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.size = Pt(10)

# ============================================================
# PORTADA
# ============================================================
doc.add_paragraph('')
doc.add_paragraph('')
T('SISTEMA DE CURSOS VIRTUALES', 0)
T('Guia completa del codigo: de principio a fin', 2)
P('')
P('Una explicacion facil de entender con ejemplos practicos')
P('')
p = doc.add_paragraph()
r = p.add_run('Hecho con Laravel 12 | PHP 8.2 | SQLite | Tailwind CSS')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

# ============================================================
# 1. INTRODUCCION
# ============================================================
T('1. Que hace esta aplicacion?')
P('Este proyecto es un panel de administracion para gestionar cursos virtuales. Piensa en el como el sistema que usa una academia online para:')
B('Crear y administrar cursos (con precio gratuito o de pago)')
B('Registrar estudiantes')
B('Inscribir estudiantes en cursos (respetando los cupos disponibles)')
B('Organizar el contenido de cada curso en modulos (algunos gratuitos, otros premium)')
B('Gestionar solicitudes de acceso a cursos de pago')
B('Tener diferentes tipos de usuarios: administrador, editor e invitado')
P('')
P('Todo el sistema esta en espanol y construido con Laravel 12, el framework de PHP mas popular.')

E('Flujo completo del sistema',
  '1. Un profesor (admin) crea un curso con 30 cupos y precio S/50\n2. El profesor agrega 5 modulos: 2 gratuitos y 3 premium\n3. Un estudiante se registra y se inscribe al curso (si hay cupos)\n4. El estudiante ve los 2 modulos gratuitos\n5. Para ver los premium, solicita acceso\n6. El administrador revisa y aprueba la solicitud\n7. El estudiante ya puede ver todos los modulos',
  'Este es el ciclo de vida completo de un estudiante en el sistema.')
doc.add_page_break()

# ============================================================
# 2. TECNOLOGIAS
# ============================================================
T('2. Tecnologias utilizadas')
t = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
t.cell(0, 0).text = 'Tecnologia'
t.cell(0, 1).text = 'Para que se usa?'
t.cell(1, 0).text = 'Laravel 12'
t.cell(1, 1).text = 'El corazon de la app: maneja rutas, controladores, BD, autenticacion'
t.cell(2, 0).text = 'PHP 8.2'
t.cell(2, 1).text = 'El lenguaje de programacion principal'
t.cell(3, 0).text = 'SQLite'
t.cell(3, 1).text = 'Base de datos en un solo archivo (no necesita instalacion)'
t.cell(4, 0).text = 'Blade'
t.cell(4, 1).text = 'Sistema de plantillas de Laravel para las paginas HTML'
t.cell(5, 0).text = 'Tailwind CSS'
t.cell(5, 1).text = 'Framework de CSS para darle estilo a las paginas'
t.cell(6, 0).text = 'Vite'
t.cell(6, 1).text = 'Herramienta que compila el CSS y JavaScript'
doc.add_page_break()

# ============================================================
# 3. ESTRUCTURA
# ============================================================
T('3. Estructura del proyecto')
P('Asi esta organizado el codigo. Cada carpeta tiene un proposito especifico:')
C('''
cursos_virtuales2/
 app/            Codigo PHP (controladores, modelos, middleware)
 routes/         Rutas web (web.php)
 resources/      Vistas Blade + CSS + JavaScript
 database/       Migraciones y datos de prueba
 public/         Punto de entrada (index.php)
 config/         Configuracion de Laravel
 storage/        Archivos subidos, logs, cache
 vendor/         Librerias externas (no tocar)
'''.strip())
N('Las carpetas app/, routes/, resources/ y database/ son las mas importantes para entender el sistema.')
doc.add_page_break()

# ============================================================
# 4. FLUJO DE PETICION
# ============================================================
T('4. Flujo de una peticion (paso a paso)')
P('Cuando un usuario escribe una URL en el navegador, esto es lo que sucede internamente:')
C('''
Paso 1: Navegador pide: GET /cursos

Paso 2: public/index.php carga Laravel
        -> bootstrap/app.php inicializa todo

Paso 3: routes/web.php busca la ruta /cursos
        -> Encuentra: Route::get("/cursos", [CursoController::class, "index"])

Paso 4: Middleware "auth" verifica: ¿el usuario inicio sesion?
        -> Si NO -> redirige a /login
        -> Si SI -> deja pasar

Paso 5: CursoController@index ejecuta:
        -> $cursos = Curso::all();  (SELECT * FROM cursos)

Paso 6: Vista cursos/index.blade.php toma los cursos
        -> Renderiza una tabla HTML con todos los cursos

Paso 7: El usuario ve la lista de cursos en su navegador
'''.strip())
P('')
P('Cada vez que haces clic en un enlace o envias un formulario, este ciclo se repite.')
E('El viaje de los datos al crear un curso',
  '1. Llenas el formulario en /cursos/create y haces clic en Guardar\n2. El navegador envia POST /cursos con los datos\n3. Laravel ejecuta CursoController@store\n4. Valida: nombre obligatorio, cupos debe ser numero >= 1\n5. Si todo OK -> Curso::create() guarda en SQLite\n6. Redirige a /cursos con mensaje "Curso creado."',
  'Cada operacion de crear/editar/eliminar sigue exactamente este patron.')
doc.add_page_break()

# ============================================================
# 5. RUTAS
# ============================================================
T('5. Las rutas (routes/web.php)')
P('El archivo de rutas es como el mapa de la aplicacion. Le dice a Laravel que hacer cuando alguien visita cada URL.')

T('5.1 Autenticacion', 2)
C("Route::get('/login', [AuthController::class, 'showLogin']);  // Muestra formulario login")
C("Route::post('/login', [AuthController::class, 'login']);     // Procesa inicio sesion")
E('Ruta de login',
  'Tu escribes: http://localhost/login\nGET -> Laravel llama a AuthController->showLogin()\n   -> Muestra la vista auth/login.blade.php\n   -> Ves un formulario con email y contrasena\nPOST -> Envia email + password\n   -> AuthController->login() verifica las credenciales',
  'GET es para ver paginas, POST es para enviar datos del formulario.')

T('5.2 Rutas protegidas (requieren inicio de sesion)', 2)
C('''
Route::middleware(['auth', 'nocache'])->group(function () {
    // Todas estas rutas requieren autenticacion
    Route::get('/dashboard', [AuthController::class, 'dashboard']);
    Route::post('/logout', [AuthController::class, 'logout']);
    // ... y todas las demas rutas del sistema
});
'''.strip())
N("auth verifica que el usuario haya iniciado sesion. nocache evita que el navegador guarde paginas en cache por seguridad.")

T('5.3 CRUD de cursos (7 rutas)', 2)
C('''
Route::get("/cursos", [CursoController::class, "index"]);       // Listar cursos
Route::get("/cursos/create", [CursoController::class, "create"]); // Formulario crear
Route::post("/cursos", [CursoController::class, "store"]);        // Guardar nuevo
Route::get("/cursos/{curso}", [CursoController::class, "show"]);  // Ver detalle
Route::get("/cursos/{curso}/edit", [CursoController::class, "edit"]); // Form editar
Route::put("/cursos/{curso}", [CursoController::class, "update"]);    // Actualizar
Route::delete("/cursos/{curso}", [CursoController::class, "destroy"]);// Eliminar
'''.strip())
E('Lectura de rutas con parametros',
  'Si visitas: /cursos/5/edit\nLaravel sabe que {curso} = 5\nEntonces busca el curso con ID=5 y lo pasa al controlador\nSi el curso no existe -> Laravel da error 404 automaticamente',
  'El {curso} entre llaves es un parametro que Laravel extrae de la URL y convierte en objeto (Route Model Binding).')

T('5.4 Rutas con roles (solo administradores)', 2)
C('''
Route::middleware("role:administrador")->group(function () {
    Route::get("/users", [UserController::class, "index"]);
    Route::get("/users/{user}/edit", [UserController::class, "edit"]);
    Route::put("/users/{user}", [UserController::class, "update"]);
});
'''.strip())
E('Que pasa si un editor entra a /users?',
  'El middleware role:administrador revisa:\n - El usuario tiene role = "administrador"?\n - Si NO -> aparece pagina: 403 No tienes permiso\n - Si SI -> puede ver la lista de usuarios',
  'Los roles protegen las rutas: solo administradores pueden gestionar usuarios.')

T('5.5 Resumen de todas las rutas', 2)
routes_table = doc.add_table(rows=18, cols=3, style='Light Shading Accent 1')
routes_table.cell(0, 0).text = 'URL'
routes_table.cell(0, 1).text = 'Metodo'
routes_table.cell(0, 2).text = 'Que hace?'
routes_data = [
    ('/login', 'GET', 'Muestra formulario de inicio de sesion'),
    ('/login', 'POST', 'Inicia sesion (valida email y password)'),
    ('/dashboard', 'GET', 'Muestra el panel principal'),
    ('/logout', 'POST', 'Cierra la sesion'),
    ('/cursos', 'GET', 'Lista todos los cursos'),
    ('/cursos/create', 'GET', 'Muestra formulario para crear curso'),
    ('/cursos/{id}', 'GET', 'Muestra detalle de un curso'),
    ('/cursos/{id}/edit', 'GET', 'Formulario para editar curso'),
    ('/estudiantes', 'GET', 'Lista estudiantes'),
    ('/inscripciones', 'GET', 'Lista inscripciones'),
    ('/cursos/{id}/modulos', 'GET', 'Lista modulos de un curso'),
    ('/users', 'GET', '[Admin] Lista usuarios del sistema'),
    ('/accesos', 'GET', '[Admin] Lista solicitudes de acceso'),
    ('/cursos/{id}/solicitar-acceso', 'POST', 'Solicita acceso a un curso'),
    ('/cursos/{id}/modulos/{id}/ver', 'GET', 'Ver contenido de un modulo'),
    ('/cursos/{id}/modulos/crear', 'GET', 'Crear nuevo modulo'),
    ('/cursos/{id}/modulos/{id}/editar', 'GET', 'Editar modulo'),
]
for i, (url, method, desc) in enumerate(routes_data):
    routes_table.cell(i+1, 0).text = url
    routes_table.cell(i+1, 1).text = method
    routes_table.cell(i+1, 2).text = desc
doc.add_page_break()

# ============================================================
# 6. CONTROLADORES
# ============================================================
T('6. Los controladores (la logica de negocio)')
P('Los controladores son el cerebro de la aplicacion. Cada controlador agrupa funciones relacionadas.')

T('6.1 AuthController - Inicio de sesion', 2)
C('''
class AuthController extends Controller
{
    // Muestra el formulario de login
    public function showLogin()
    {
        return view("auth.login");
    }

    // Procesa el inicio de sesion
    public function login(Request $request)
    {
        // 1. Valida que email y password existan
        $credentials = $request->validate([
            "email"    => ["required", "email"],
            "password" => ["required"]
        ]);

        // 2. Intenta autenticar (busca el usuario en la BD)
        if (Auth::attempt($credentials)) {
            $request->session()->regenerate(); // seguridad
            return redirect()->route("dashboard")
                ->with("success", "Bienvenido al sistema");
        }

        // 3. Si no coincide -> mensaje de error
        return back()->withErrors([
            "email" => "Las credenciales no son correctas.",
        ]);
    }
}
'''.strip())

E('Que pasa cuando alguien inicia sesion?',
  'Usuario escribe: email=juan@mail.com, password=123456\n\n1. Laravel busca en la tabla "users" un email = juan@mail.com\n2. Toma el password guardado (esta encriptado: $2y$10$...)\n3. Compara: ?123456 encriptado es igual al guardado?\n4. Si SI -> inicia la sesion, guarda los datos del usuario\n5. Si NO -> mensaje "Las credenciales no son correctas"',
  'Nunca se guardan las contrasenas en texto plano. Laravel las encripta automaticamente con bcrypt.')

T('6.2 CursoController - Gestion de cursos', 2)
C('''
class CursoController extends Controller
{
    public function index()
    {
        $cursos = Curso::all(); // SELECT * FROM cursos
        return view("cursos.index", compact("cursos"));
    }

    public function store(Request $request)
    {
        // Validar los datos del formulario
        $request->validate([
            "nombre" => "required",
            "cupos"  => "required|integer|min:1",
            "precio" => "nullable|numeric|min:0",
        ]);

        // Si el precio es 0 o vacio -> es gratis
        $data = $request->all();
        $data["es_gratis"] = ($request->precio == 0 || $request->precio == null);

        Curso::create($data); // INSERT INTO cursos
        return redirect()->route("cursos.index")
            ->with("success", "Curso creado.");
    }

    public function destroy(Curso $curso)
    {
        $curso->delete(); // DELETE FROM cursos WHERE id = ?
        return redirect()->route("cursos.index")
            ->with("success", "Curso eliminado.");
    }
}
'''.strip())

E('Creacion de un curso paso a paso',
  'Formulario: nombre=Matematicas, cupos=30, precio=50\n\n1. store() recibe los datos\n2. Validacion: nombre OK, cupos = 30 (numero >= 1), precio = 50 (numero)\n3. Como precio=50 (no es 0), es_gratis = false\n4. Curso::create() guarda en SQLite:\n   INSERT INTO cursos (nombre, cupos, precio, es_gratis)\n   VALUES ("Matematicas", 30, 50.00, 0)\n5. Redirige a /cursos con mensaje "Curso creado."',
  'Siempre se valida antes de guardar. Esto evita datos incorrectos.')

E('Que hace "Curso $curso" en el parametro?',
  'Ruta: /cursos/{curso}/edit\nURL: /cursos/5/edit\n\nLaravel ve que {curso}=5 y automaticamente hace:\n  $curso = Curso::find(5);  // SELECT * FROM cursos WHERE id=5\nSi el ID no existe -> error 404.\nEsto se llama Route Model Binding y evita escribir la consulta manualmente.',
  'El parametro de la ruta se convierte automaticamente en el modelo.')

T('6.3 EstudianteController - Gestion de estudiantes', 2)
C('''
class EstudianteController extends Controller
{
    public function index()
    {
        // Trae todos los estudiantes con sus cursos
        $estudiantes = Estudiante::with("cursos")->get();
        return view("estudiantes.index", compact("estudiantes"));
    }

    public function store(Request $request)
    {
        $request->validate([
            "nombre" => "required",
            "email"  => "required|email|unique:estudiantes",
        ]);
        Estudiante::create($request->all());
        return redirect()->route("estudiantes.index")
            ->with("success", "Estudiante registrado.");
    }
}
'''.strip())

E('La validacion "unique"',
  'Si dos personas se registran con el mismo email:\n\n1er registro: juan@gmail.com -> OK, guarda\n2do registro: juan@gmail.com -> ERROR: "El email ya esta registrado"\n\nEsto evita duplicados en la tabla de estudiantes.',
  'La regla unique:estudiantes revisa que el email no exista ya en esa tabla.')

T('6.4 InscripcionController - Inscripciones (la mas interesante)', 2)
C('''
public function store(Request $request)
{
    // 1. Validar que los IDs existan en sus tablas
    $request->validate([
        "curso_id"      => "required|exists:cursos,id",
        "estudiante_id" => "required|exists:estudiantes,id",
    ]);

    // 2. Buscar el curso
    $curso = Curso::find($request->curso_id);

    // 3. Verificar si HAY CUPOS disponibles
    if ($curso->cuposDisponibles() <= 0) {
        return back()->with("error", "No hay cupos disponibles.");
    }

    // 4. Verificar que NO este ya inscrito
    $existe = Inscripcion::where("curso_id", $request->curso_id)
        ->where("estudiante_id", $request->estudiante_id)
        ->exists();

    if ($existe) {
        return back()->with("error", "El estudiante ya esta inscrito.");
    }

    // 5. Si todo OK -> inscribir
    Inscripcion::create($request->all());
    return redirect()->route("inscripciones.index")
        ->with("success", "Inscripcion realizada.");
}
'''.strip())

E('Que pasa cuando un estudiante se inscribe a un curso lleno?',
  'Curso: Matematicas con cupos=30 y ya hay 30 inscritos\n\n1. Estudiante intenta inscribirse\n2. $curso->cuposDisponibles() calcula: 30 - 30 = 0\n3. 0 <= 0 -> es TRUE -> muestra error: "No hay cupos disponibles"\n4. El estudiante NO puede inscribirse',
  'El metodo cuposDisponibles() resta los cupos totales menos los inscritos actuales.')

E('Que pasa si intentas inscribir al mismo estudiante dos veces?',
  '1er intento: Maria se inscribe a Matematicas -> OK\n2do intento: Maria se inscribe a Matematicas otra vez\n\nLa consulta busca:\n  ?Existe una inscripcion con curso_id=1 Y estudiante_id=3?\n  -> SI, ya existe\n  -> Muestra error: "El estudiante ya esta inscrito en este curso."',
  'Se verifica duplicado combinando curso_id + estudiante_id antes de insertar.')

T('6.5 ModuloController - Modulos de los cursos', 2)
C('''
public function ver(Curso $curso, Modulo $modulo)
{
    // ?El usuario actual es un estudiante registrado?
    $estudiante = Estudiante::where("email", auth()->user()->email)->first();

    // Si el modulo NO es de prueba -> verificar acceso
    if (!$modulo->es_prueba) {
        $tieneAcceso = false;

        if ($estudiante) {
            // ?Tiene un acceso aprobado para este curso?
            $acceso = Acceso::where("curso_id", $curso->id)
                ->where("estudiante_id", $estudiante->id)
                ->where("estado", "aprobado")
                ->first();
            $tieneAcceso = $acceso !== null;
        }

        // Si no tiene acceso y no es admin -> bloqueado
        if (!$tieneAcceso && auth()->user()->role !== "administrador") {
            return redirect()->route("cursos.modulos.index", $curso)
                ->with("error", "Necesitas acceso completo.");
        }
    }

    // Si es modulo de prueba, o tiene acceso, o es admin -> puede verlo
    return view("modulos.ver", compact("curso", "modulo"));
}
'''.strip())

E('Quien puede ver cada tipo de modulo?',
  'Curso: Ingles Basico\n  Modulo 1: "El alfabeto" (es_prueba = true)  -> LO VE TODO EL MUNDO!\n  Modulo 2: "Verbos" (es_prueba = false) -> SOLO con acceso aprobado\n  Modulo 3: "Conversacion" (es_prueba = false) -> SOLO con acceso aprobado\n\nSi un estudiante SIN acceso intenta ver el Modulo 3:\n  -> "Necesitas acceso completo para ver este modulo."\n\nSi un ADMINISTRADOR intenta verlo:\n  -> Lo ve sin restricciones (el admin tiene pase libre)',
  'El sistema distingue entre modulos de prueba (gratis para todos) y premium (requieren aprobacion).')

T('6.6 AccesoController - Solicitudes de acceso', 2)
C('''
// El estudiante solicita acceso a un curso de pago
public function solicitar(Curso $curso)
{
    $estudiante = Estudiante::where("email", auth()->user()->email)->first();

    if (!$estudiante) {
        return back()->with("error", "No eres estudiante.");
    }

    // ?Ya solicito acceso antes?
    $yaExiste = Acceso::where("curso_id", $curso->id)
        ->where("estudiante_id", $estudiante->id)->exists();

    if ($yaExiste) {
        return back()->with("error", "Ya tienes una solicitud.");
    }

    Acceso::create([
        "curso_id"        => $curso->id,
        "estudiante_id"   => $estudiante->id,
        "estado"          => "pendiente",
        "fecha_solicitud" => now(),
    ]);
    return back()->with("success", "Solicitud enviada.");
}

// Admin aprueba la solicitud
public function aprobar(Acceso $acceso)
{
    $acceso->update([
        "estado"           => "aprobado",
        "fecha_aprobacion" => now(),
    ]);
    return back()->with("success", "Acceso aprobado.");
}

// Admin rechaza la solicitud
public function rechazar(Acceso $acceso)
{
    $acceso->update(["estado" => "rechazado"]);
    return back()->with("success", "Acceso rechazado.");
}
'''.strip())

E('El ciclo completo de una solicitud de acceso',
  '1. Estudiante ve un curso de pago y hace clic en "Solicitar acceso"\n2. Se crea un registro en "accesos" con estado=pendiente\n3. El administrador entra a /accesos y ve:\n     Estudiante: Juan Perez\n     Curso: Ingles Basico (S/50)\n     Estado: Pendiente  [Aprobar] [Rechazar]\n4. Admin hace clic en [Aprobar]\n5. El estado cambia a "aprobado" y se guarda la fecha\n6. Ahora el estudiante puede ver todos los modulos del curso',
  'El flujo de aprobacion tiene 3 estados: pendiente -> aprobado/rechazado.')

T('6.7 UserController - Gestion de usuarios (solo admin)', 2)
C('''
public function update(Request $request, User $user)
{
    $request->validate([
        "name"  => "required|string|max:255",
        "email" => "required|email|unique:users,email," . $user->id,
        "role"  => "required|in:administrador,editor,invitado",
        "photo" => "nullable|image|mimes:jpg,jpeg,png,webp|max:2048",
    ]);

    $data = $request->only("name", "email", "role");

    // Si subio una foto nueva
    if ($request->hasFile("photo")) {
        // Borrar la foto anterior si existe
        if ($user->photo && Storage::disk("public")->exists($user->photo)) {
            Storage::disk("public")->delete($user->photo);
        }
        // Guardar la nueva foto
        $data["photo"] = $request->file("photo")->store("users", "public");
    }

    $user->update($data);
    return redirect()->route("users.index")
        ->with("success", "Usuario actualizado.");
}
'''.strip())

E('Subida de foto de perfil',
  '1. Admin selecciona una foto "juan.jpg" (maximo 2MB)\n2. Laravel valida: ?es imagen? ?es jpg/png/webp? ?pesa <= 2048KB?\n3. Si todo OK -> guarda en storage/app/public/users/juan.jpg\n4. Borra la foto anterior si existia\n5. Guarda la ruta "users/juan.jpg" en la base de datos\n6. La foto se muestra en la lista de usuarios',
  'Las fotos se guardan en el disco publico y se accede a traves de un enlace simbolico.')

doc.add_page_break()

# ============================================================
# 7. MODELOS
# ============================================================
T('7. Los modelos (la estructura de los datos)')
P('Los modelos en Laravel representan las tablas de la base de datos. Piensa en ellos como la plantilla que define que datos guardamos y como se relacionan entre si.')

T('7.1 Modelo Curso', 2)
C('''
class Curso extends Model
{
    // Campos que se pueden llenar desde un formulario
    protected $fillable = ["nombre", "descripcion", "cupos", "precio", "es_gratis"];

    // Relaciones (como se conecta con otras tablas):

    // Un curso tiene muchas inscripciones
    public function inscripciones() {
        return $this->hasMany(Inscripcion::class);
    }

    // Un curso tiene muchos estudiantes (a traves de inscripciones)
    public function estudiantes() {
        return $this->belongsToMany(Estudiante::class, "inscripciones");
    }

    // Un curso tiene muchos modulos, ordenados por "orden"
    public function modulos() {
        return $this->hasMany(Modulo::class)->orderBy("orden");
    }

    // Calcula: cupos totales - inscritos actuales
    public function cuposDisponibles() {
        return $this->cupos - $this->inscripciones()->count();
    }
}
'''.strip())

E('Usando las relaciones del modelo',
  '$curso = Curso::find(1);\n\necho $curso->nombre;           // "Matematicas"\necho $curso->descripcion;       // "Curso de algebra basica"\necho $curso->cupos;             // 30\necho $curso->cuposDisponibles(); // 25 (30 - 5 inscritos)\n\nforeach ($curso->modulos as $modulo) {\n    echo $modulo->titulo;  // "Introduccion", "Sumas", "Restas"...\n}\n\nforeach ($curso->estudiantes as $estudiante) {\n    echo $estudiante->nombre;  // "Juan", "Maria", "Pedro"...\n}',
  'Las relaciones te permiten navegar entre tablas como si fueran propiedades del objeto.')

T('7.2 Todos los modelos y sus relaciones', 2)
models_table = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
models_table.cell(0, 0).text = 'Modelo'
models_table.cell(0, 1).text = 'Relaciones'
models_data = [
    ('Curso', '-> tiene muchos Modulos\n-> tiene muchas Inscripciones\n-> tiene muchos Estudiantes (a traves de Inscripciones)\n-> tiene muchas Solicitudes de Acceso'),
    ('Estudiante', '-> esta en muchos Cursos (a traves de Inscripciones)\n-> tiene muchas Solicitudes de Acceso'),
    ('Inscripcion', '-> pertenece a un Curso\n-> pertenece a un Estudiante'),
    ('Modulo', '-> pertenece a un Curso'),
    ('Acceso', '-> pertenece a un Curso\n-> pertenece a un Estudiante'),
    ('User', '(no tiene relaciones personalizadas - es el usuario del sistema)'),
]
for i, (model, rel) in enumerate(models_data):
    models_table.cell(i+1, 0).text = model
    models_table.cell(i+1, 1).text = rel

E('Diagrama de relaciones entre tablas',
  '  Curso ----1---a---n----> Inscripcion <---n---a---1---- Estudiante\n    |                                              |\n    |1                                             |1\n    |                                              |\n    v                                              v\n  Modulo                                        Acceso\n\nExplicacion:\n- 1 Curso -> tiene muchos (n) Modulos\n- 1 Curso -> tiene muchas (n) Inscripciones\n- 1 Estudiante -> tiene muchas (n) Inscripciones\n- 1 Estudiante -> tiene muchas (n) Solicitudes de Acceso\n- La tabla inscripciones es una tabla puente (muchos-a-muchos)',
  'Las flechas indican UNO a MUCHOS. Una tabla puente como inscripciones conecta cursos con estudiantes.')

doc.add_page_break()

# ============================================================
# 8. MIDDLEWARE
# ============================================================
T('8. Los middleware (los filtros de seguridad)')
P('Los middleware son como guardias que revisan cada peticion antes de que llegue al controlador. Piensa en ellos como filtros de seguridad.')

T('8.1 VerifyAuth - El usuario inicio sesion?', 2)
C('''
class VerifyAuth
{
    public function handle(Request $request, Closure $next)
    {
        if (!Auth::check()) {  // ?NO hay sesion iniciada?
            if ($request->expectsJson()) {
                return response()->json(["error" => "No autenticado"], 401);
            }
            return redirect()->route("login.form");  // -> al login
        }
        return $next($request);  // -> adelante, pasa
    }
}
'''.strip())

E('Que pasa si un usuario no autenticado entra a /cursos?',
  '1. El middleware VerifyAuth revisa: ?hay sesion iniciada?\n2. NO -> redirige a /login\n3. El usuario ve el formulario de inicio de sesion\n\nAsi se protegen TODAS las rutas del sistema (excepto /login).',
  'Sin este middleware, cualquiera podria acceder a cualquier pagina.')

T('8.2 RoleMiddleware - Tiene el rol adecuado?', 2)
C('''
class RoleMiddleware
{
    public function handle(Request $request, Closure $next, ...$roles)
    {
        if (!auth()->check()) {
            return redirect()->route("login.form");
        }
        if (!in_array(auth()->user()->role, $roles)) {
            abort(403, "No tienes permiso para realizar esta accion.");
        }
        return $next($request);
    }
}
'''.strip())

E('Que pasa si un invitado intenta eliminar un curso?',
  'Ruta: DELETE /cursos/5 (protegida con middleware)\n\n1. El usuario "invitado" hace clic en Eliminar\n2. RoleMiddleware revisa: ?el usuario es administrador?\n3. El usuario tiene role=invitado\n4. "invitado" NO esta en la lista de roles permitidos\n5. -> Aparece: 403 No tienes permiso para realizar esta accion.',
  'Los roles evitan que usuarios no autorizados hagan acciones peligrosas.')

T('8.3 NoCacheMiddleware - Sin cache por seguridad', 2)
C('''
class NoCacheMiddleware
{
    public function handle(Request $request, Closure $next)
    {
        $response = $next($request);
        return $response->withHeaders([
            "Cache-Control" => "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma"        => "no-cache",
            "Expires"       => "0",
        ]);
    }
}
'''.strip())

N("Este middleware evita que el navegador guarde paginas en cache. Asi, si un usuario cierra sesion y alguien usa su computadora, no podra ver las paginas anteriores desde el historial.")

T('Resumen de los 3 middleware', 2)
mw_table = doc.add_table(rows=4, cols=3, style='Light Shading Accent 1')
mw_table.cell(0, 0).text = 'Middleware'
mw_table.cell(0, 1).text = 'Que revisa?'
mw_table.cell(0, 2).text = 'Que pasa si falla?'
mw_data = [
    ('auth (VerifyAuth)', 'El usuario inicio sesion?', 'Redirige al login'),
    ('role (RoleMiddleware)', 'Tiene el rol adecuado?', 'Error 403 (prohibido)'),
    ('nocache (NoCacheMiddleware)', 'Agrega headers anti-cache', '(no falla, solo agrega seguridad)'),
]
for i, (a, b, c) in enumerate(mw_data):
    mw_table.cell(i+1, 0).text = a
    mw_table.cell(i+1, 1).text = b
    mw_table.cell(i+1, 2).text = c

doc.add_page_break()

# ============================================================
# 9. VISTAS
# ============================================================
T('9. Las vistas (Blade - el HTML)')
P('Las vistas son los archivos que generan el HTML que ves en el navegador. Laravel usa Blade, un sistema de plantillas que permite mezclar PHP con HTML de forma sencilla.')

T('9.1 El layout principal', 2)
C('''
<html>
<head><title>@yield("title", "Cursos Virtuales")</title></head>
<body>
  <nav>
    <a href="{{ route("cursos.index") }}">Cursos</a>
    <a href="{{ route("estudiantes.index") }}">Estudiantes</a>
    <a href="{{ route("inscripciones.index") }}">Inscripciones</a>

    @if(auth()->user()->role === "administrador")
      <a href="{{ route("users.index") }}">Usuarios</a>
      <a href="{{ route("accesos.index") }}">Accesos</a>
    @endif

    <span>{{ auth()->user()->name }}</span>
    <form method="POST" action="{{ route("logout") }}">
      @csrf
      <button>Cerrar sesion</button>
    </form>
  </nav>

  <div class="container">
    @yield("content")  <!-- Aqui se inserta el contenido de cada pagina -->
  </div>
</body>
</html>
'''.strip())

E('Como funciona @yield y @section?',
  'En el layout:  <div>@yield("content")</div>\n\nEn cada pagina (ej: cursos/index.blade.php):\n  @extends("layouts.app")       <- usa el layout\n  @section("title", "Lista")    <- llena el titulo\n  @section("content")           <- llena el contenido\n      <h1>Lista de Cursos</h1>\n      <table>...</table>\n  @endsection\n\nResultado final en el navegador:\n  <html>\n    <head><title>Lista</title></head>\n    <body>\n      <nav>...</nav>\n      <div class="container">\n        <h1>Lista de Cursos</h1>\n        <table>...</table>\n      </div>\n    </body>\n  </html>',
  'El layout es el molde y cada pagina solo llena las secciones variables.')

T('9.2 Vista de lista de cursos', 2)
C('''
@foreach($cursos as $curso)
  <tr>
    <td>{{ $curso->nombre }}</td>
    <td>{{ $curso->cupos }}</td>
    <td>{{ $curso->cuposDisponibles() }}</td>
    <td>
      @if($curso->es_gratis)
        <span>GRATIS</span>
      @else
        S/ {{ number_format($curso->precio, 2) }}
      @endif
    </td>
    <td>
      <a href="{{ route("cursos.edit", $curso) }}">Editar</a>
      @if(auth()->user()->role === "administrador")
        <form method="POST" action="{{ route("cursos.destroy", $curso) }}">
          @csrf @method("DELETE")
          <button>Eliminar</button>
        </form>
      @endif
    </td>
  </tr>
@endforeach
'''.strip())

E('Como se muestra el precio?',
  '@if($curso->es_gratis)         <- Si es gratis\n    <span>GRATIS</span>          <- muestra etiqueta\n@else                             <- Si no\n    S/ {{ number_format($curso->precio, 2) }}  <- "S/ 50.00"\n@endif\n\nResultado en la tabla:\n  Matematicas | 30 | 25 | S/ 50.00 |\n  Historia    | 20 | 20 | GRATIS   |',
  'La vista decide como mostrar los datos segun su valor.')

T('9.3 Vista de login', 2)
C('''
<form method="POST" action="{{ route("login") }}">
    @csrf  <!-- Token de seguridad obligatorio -->

    <label>Email</label>
    <input type="email" name="email" value="{{ old("email") }}">
    @error("email")
        <div>{{ $message }}</div>  <!-- Mensaje de error si la validacion fallo -->
    @enderror

    <label>Contrasena</label>
    <input type="password" name="password">

    <button type="submit">Ingresar</button>
</form>
'''.strip())

N("@csrf genera un token de seguridad unico. Sin el, Laravel rechaza el formulario. @error('email') muestra el mensaje de error si la validacion fallo.")

doc.add_page_break()

# ============================================================
# 10. BASE DE DATOS
# ============================================================
T('10. La base de datos (migraciones y tablas)')
P('Las migraciones son archivos que describen la estructura de las tablas. Piensa en ellas como el plano de construccion de la base de datos.')

db_table = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
db_table.cell(0, 0).text = 'Tabla'
db_table.cell(0, 1).text = 'Columnas principales'
db_table.cell(0, 2).text = 'Notas'
db_data = [
    ('cursos', 'id, nombre, descripcion, cupos, precio, es_gratis', 'precio=0 -> es_gratis=true'),
    ('estudiantes', 'id, nombre, email, telefono', 'email UNICO'),
    ('inscripciones', 'id, curso_id, estudiante_id', 'Tabla puente, FK a cursos y estudiantes'),
    ('modulos', 'id, curso_id, titulo, contenido, orden, es_prueba', 'es_prueba=true = modulo gratuito'),
    ('accesos', 'id, estudiante_id, curso_id, estado, fecha_solicitud, fecha_aprobacion', 'estado: pendiente/aprobado/rechazado'),
    ('users', 'id, name, email, password, role, photo', 'role: administrador/editor/invitado'),
]
for i, (a, b, c) in enumerate(db_data):
    db_table.cell(i+1, 0).text = a
    db_table.cell(i+1, 1).text = b
    db_table.cell(i+1, 2).text = c

P('')
E('Como se relacionan los datos entre tablas? Ejemplo con datos reales',
  'Tabla cursos:\n  id=1, nombre=Matematicas, cupos=30, precio=50.00\n\nTabla estudiantes:\n  id=1, nombre=Juan Perez, email=juan@mail.com\n  id=2, nombre=Maria Lopez, email=maria@mail.com\n\nTabla inscripciones:\n  curso_id=1, estudiante_id=1  (Juan esta en Matematicas)\n  curso_id=1, estudiante_id=2  (Maria esta en Matematicas)\n\nTabla modulos:\n  curso_id=1, titulo=Sumas, orden=1, es_prueba=true\n  curso_id=1, titulo=Derivadas, orden=3, es_prueba=false\n\nTabla accesos:\n  estudiante_id=1, curso_id=1, estado=aprobado\n  (Juan tiene acceso a los modulos premium de Matematicas)',
  'Asi se ven los datos en la vida real. Cada tabla guarda una parte de la informacion.')

doc.add_page_break()

# ============================================================
# 11. DIAGRAMA DE FLUJO COMPLETO
# ============================================================
T('11. Diagrama de flujo completo del sistema')
P('Aqui tienes el recorrido COMPLETO que hace un estudiante desde que llega al sistema hasta que accede a un modulo premium:')
C('''
1. USUARIO ESCRIBE: http://localhost:8000
   -> Redirige a /login

2. FORMULARIO DE LOGIN
   Email: juan@mail.com
   Password: ******
   -> AuthController::login()
   -> Auth::attempt() busca en tabla "users"
   -> Si OK -> redirige al dashboard
   -> Si falla -> mensaje "Credenciales incorrectas"

3. DASHBOARD (panel principal)
   Bienvenido, Juan
   [Cursos]  [Estudiantes]  [Inscripciones]
   (si eres admin tambien: [Usuarios] [Accesos])

4. LISTA DE CURSOS (/cursos)
   Curso       Cupos  Disponibles  Precio    Acciones
   Matematicas   30       25       S/50.00  [Modulos]
   Historia      20       20       GRATIS   [Modulos]

5. MODULOS DEL CURSO (/cursos/1/modulos)
   Orden  Titulo         Tipo            Accion
    1     Sumas          (gratis)        [Ver]
    2     Derivadas      (premium)       [Solicitar acceso]

   -> Juan puede ver "Sumas" (es gratuito)
   -> Para "Derivadas" necesita hacer clic en "Solicitar acceso"

6. SOLICITAR ACCESO
   POST /cursos/1/solicitar-acceso
   -> AccesoController::solicitar()
   -> Se crea registro: estado=pendiente
   -> Mensaje: "Solicitud enviada. El admin revisara."

7. ADMIN APRUEBA (/accesos)
   Estudiante  Curso       Estado     Accion
   Juan Perez  Matematicas Pendiente  [Aprobar]

   Admin hace clic en [Aprobar]
   -> AccesoController::aprobar()
   -> estado cambia a "aprobado"
   -> Mensaje: "Acceso aprobado correctamente."

8. JUAN VE MODULOS PREMIUM
   Vuelve a /cursos/1/modulos
   Ahora puede hacer clic en [Ver] en "Derivadas"
   -> ModuloController::ver()
   -> Verifica: ?tiene acceso aprobado? -> SI
   -> Muestra el contenido del modulo

============================================
    Este es el ciclo completo del sistema.
============================================
'''.strip())

doc.add_page_break()

# ============================================================
# 12. GLOSARIO
# ============================================================
T('12. Glosario de terminos')
gl_table = doc.add_table(rows=14, cols=2, style='Light Shading Accent 1')
gl_table.cell(0, 0).text = 'Termino'
gl_table.cell(0, 1).text = 'Significado'
glossary = [
    ('Ruta (Route)', 'Define que URL lleva a que controlador. Ej: /cursos -> CursoController'),
    ('Controlador', 'La clase PHP que contiene la logica para cada pagina'),
    ('Modelo', 'Representa una tabla de la BD. Ej: Curso.php = tabla cursos'),
    ('Vista (Blade)', 'Archivo HTML con PHP que genera lo que ves en el navegador'),
    ('Middleware', 'Filtro que revisa las peticiones (?esta autenticado? ?es admin?)'),
    ('Migracion', 'Archivo que crea o modifica una tabla en la base de datos'),
    ('Validacion', 'Reglas que verifican que los datos del formulario sean correctos'),
    ('Eloquent ORM', 'Permite trabajar con la BD usando objetos PHP en lugar de SQL'),
    ('Route Model Binding', 'Laravel convierte automaticamente un ID de URL en el modelo'),
    ('Blade @yield', 'Espacio en el layout donde se inserta el contenido de cada pagina'),
    ('@csrf', 'Token de seguridad que Laravel exige en todos los formularios'),
    ('SQLite', 'Base de datos en un solo archivo (database.sqlite)'),
    ('Rol (role)', 'Permiso del usuario: administrador, editor o invitado'),
]
for i, (term, meaning) in enumerate(glossary):
    gl_table.cell(i+1, 0).text = term
    gl_table.cell(i+1, 1).text = meaning

doc.add_paragraph('')
N("Este glosario te ayudara a recordar los conceptos clave mientras exploras el codigo.")

# ============================================================
# FINAL
# ============================================================
doc.add_paragraph('')
T('Fin del documento!', 2)
P('Ahora ya conoces todo el sistema de principio a fin:')
B('Sabes como esta organizado el proyecto')
B('Entiendes el flujo de una peticion (Ruta -> Controlador -> Modelo -> Vista)')
B('Conoces cada controlador y lo que hace')
B('Comprendes las relaciones entre las tablas de la base de datos')
B('Sabes como funcionan los roles y los permisos')
B('Puedes seguir el recorrido completo de un estudiante en el sistema')

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('A programar se ha dicho!')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# GUARDAR
# ============================================================
output_path = r'C:\Users\brian\OneDrive\Desktop\cursos_virtuales2\explicacion_codigo.docx'
doc.save(output_path)
print('Documento generado exitosamente: ' + output_path)
