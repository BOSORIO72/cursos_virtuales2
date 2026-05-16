import json, sys
data = json.loads(sys.stdin.read())
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def add_para(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_example(ttitle, code_text, explanation):
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(">> EJEMPLO: " + ttitle)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
    run.font.size = Pt(11)
    add_code(code_text)
    add_para("=> " + explanation)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run("NOTA: " + text)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(10)

# PORTADA
doc.add_paragraph("")
doc.add_paragraph("")
add_title("SISTEMA DE CURSOS VIRTUALES", level=0)
add_title("Guia completa del codigo: de principio a fin", level=2)
add_para("")
add_para("Una explicacion facil de entender con ejemplos practicos")
add_para("")
p = doc.add_paragraph()
run = p.add_run("Hecho con Laravel 12 | PHP 8.2 | SQLite | Tailwind CSS")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

# 1
add_title("1. Que hace esta aplicacion?", level=1)
add_para("Este proyecto es un panel de administracion para gestionar cursos virtuales. Piensa en el como el sistema que usa una academia online para:")
add_bullet("Crear y administrar cursos (con precio gratuito o de pago)")
add_bullet("Registrar estudiantes")
add_bullet("Inscribir estudiantes en cursos (respetando los cupos disponibles)")
add_bullet("Organizar el contenido de cada curso en modulos (algunos gratuitos, otros premium)")
add_bullet("Gestionar solicitudes de acceso a cursos de pago")
add_bullet("Tener diferentes tipos de usuarios: administrador, editor e invitado")
add_para("")
add_para("Todo el sistema esta en espanol y construido con Laravel 12, el framework de PHP mas popular.")
add_example("Flujo completo del sistema","1. Un profesor (admin) crea un curso de Matematicas con 30 cupos y precio S/50\n2. El profesor agrega 5 modulos: 2 gratuitos (de prueba) y 3 premium\n3. Un estudiante se registra y se inscribe al curso (si hay cupos)\n4. El estudiante ve los 2 modulos gratuitos\n5. Para ver los premium, el estudiante solicita acceso\n6. El administrador revisa y aprueba la solicitud\n7. El estudiante ya puede ver todos los modulos","Este es el ciclo de vida completo de un estudiante en el sistema.")
doc.add_page_break()

# 2
add_title("2. Tecnologias utilizadas", level=1)
table = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
table.cell(0,0).text="Tecnologia"; table.cell(0,1).text="Para que se usa?"
table.cell(1,0).text="Laravel 12"; table.cell(1,1).text="El corazon de la app: maneja rutas, controladores, BD, autenticacion"
table.cell(2,0).text="PHP 8.2"; table.cell(2,1).text="El lenguaje de programacion en el que esta escrito todo"
table.cell(3,0).text="SQLite"; table.cell(3,1).text="Base de datos (un solo archivo, no necesita instalacion)"
table.cell(4,0).text="Blade"; table.cell(4,1).text="Sistema de plantillas de Laravel para las paginas HTML"
table.cell(5,0).text="Tailwind CSS"; table.cell(5,1).text="Framework de CSS para darle estilo a las paginas"
table.cell(6,0).text="Vite"; table.cell(6,1).text="Herramienta que compila el CSS y JavaScript para el navegador"
print("Section 2 done")
