# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

def T(text, lv=1):
    h = doc.add_heading(text, level=lv)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def C(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def P(text): doc.add_paragraph(text)
def B(text): doc.add_paragraph(text, style='List Bullet')

def E(tt, code, expl):
    doc.add_paragraph('')
    p = doc.add_paragraph(); r = p.add_run('>> EJEMPLO: ' + tt)
    r.bold = True; r.font.color.rgb = RGBColor(0x00, 0x70, 0x70); r.font.size = Pt(11)
    C(code); P('=> ' + expl)

def N(text):
    p = doc.add_paragraph(); r = p.add_run('NOTA: ' + text)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(10)
doc.add_paragraph(''); doc.add_paragraph('')
T('SISTEMA DE CURSOS VIRTUALES', 0)
T('Guia completa del codigo: de principio a fin', 2)
P(''); P('Una explicacion facil de entender con ejemplos practicos'); P('')
p = doc.add_paragraph(); r = p.add_run('Hecho con Laravel 12 | PHP 8.2 | SQLite | Tailwind CSS')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

T('1. Que hace esta aplicacion?')
P('Panel de administracion para gestionar cursos virtuales. Sirve para:')
B('Crear y administrar cursos (gratis o de pago)')
B('Registrar estudiantes')
B('Inscribir estudiantes con control de cupos')
B('Organizar contenido en modulos (gratuitos y premium)')
B('Gestionar solicitudes de acceso a cursos de pago')
P('Hecho con Laravel 12, PHP 8.2, SQLite y Tailwind CSS.')
E('Flujo completo','1. Admin crea curso\n2. Agrega modulos\n3. Estudiante se inscribe\n4. Ve modulos gratis\n5. Solicita acceso premium\n6. Admin aprueba\n7. Estudiante ve todo','Ciclo de vida completo.')

doc.add_page_break()
T('2. Tecnologias usadas')
t = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
pairs = [('Tecnologia','Uso'),('Laravel 12','Rutas, controladores, BD'),('PHP 8.2','Lenguaje principal'),('SQLite','BD archivo unico'),('Blade','Plantillas HTML'),('Tailwind CSS','Estilos'),('Vite','Compilar CSS/JS')]
for i,(a,b) in enumerate(pairs):
    t.cell(i,0).text=a; t.cell(i,1).text=b
print('sec12 done')
doc.add_page_break()
T('3. Estructura del proyecto')
C('''
app/           Codigo PHP (controladores, modelos, middleware)
routes/        Rutas (web.php)
resources/     Vistas Blade + CSS + JS
database/      Migraciones y seeders
public/        Punto de entrada (index.php)
config/        Configuracion de Laravel'''.strip())

doc.add_page_break()
T('4. Flujo de una peticion')
C('''
1. Navegador: GET /cursos
2. public/index.php carga Laravel
3. routes/web.php: CursoController@index
4. Middleware auth: verifica sesion
5. CursoController: $cursos = Curso::all()
6. Vista cursos/index.blade.php renderiza
7. Usuario ve la pagina'''.strip())

doc.add_page_break()
T('5. Las rutas (routes/web.php)')
P('El archivo de rutas es el mapa de la aplicacion.')

T('5.1 Autenticacion',2)
C("Route::get('/login', [AuthController::class, 'showLogin']);")
C("Route::post('/login', [AuthController::class, 'login']);")
E('Ruta login','GET=ver form, POST=enviar','GET muestra, POST procesa.')

T('5.2 Rutas protegidas',2)
C("Route::middleware(['auth', 'nocache'])->group(function () {")
C("    Route::get('/dashboard', [AuthController::class, 'dashboard']);")
C("    Route::post('/logout', [AuthController::class, 'logout']);")
C("});")
N('Auth verifica sesion, nocache evita cache.')

T('5.3 CRUD cursos',2)
C("Route::get('/cursos', [CursoController::class, 'index']);")
C("Route::get('/cursos/create', [CursoController::class, 'create']);")
C("Route::post('/cursos', [CursoController::class, 'store']);")
C("Route::get('/cursos/{curso}', [CursoController::class, 'show']);")
C("Route::get('/cursos/{curso}/edit', [CursoController::class, 'edit']);")
C("Route::put('/cursos/{curso}', [CursoController::class, 'update']);")
C("Route::delete('/cursos/{curso}', [CursoController::class, 'destroy']);")
E('Parametro {curso}','URL: /cursos/5/edit\nLaravel obtiene curso ID=5\nautomaticamente.','Route Model Binding: ID se convierte en objeto.')

T('5.4 Rutas solo admin',2)
C("Route::middleware('role:administrador')->group(function () {")
C("    Route::get('/users', [UserController::class, 'index']);")
C("});")
E('Sin permiso','Editor entra a /users\nRoleMiddleware: role=editor!=admin\nRespuesta: 403','Roles protegen rutas sensibles.')
print('sec345 done')
