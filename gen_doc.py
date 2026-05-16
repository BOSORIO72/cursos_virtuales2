# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

def T(text, lv=1):
    h = doc.add_heading(text, level=lv)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def C(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def P(text): doc.add_paragraph(text)
def B(text): doc.add_paragraph(text, style='List Bullet')

def E(tt, code, expl):
    doc.add_paragraph("")
    p = doc.add_paragraph(); r = p.add_run(">> EJEMPLO: " + tt)
    r.bold = True; r.font.color.rgb = RGBColor(0x00, 0x70, 0x70); r.font.size = Pt(11)
    C(code); P("=> " + expl)

def N(text):
    p = doc.add_paragraph(); r = p.add_run("NOTA: " + text)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(10)

print("Functions defined")
